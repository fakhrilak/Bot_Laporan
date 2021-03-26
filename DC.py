from docx import Document
from docx.shared import Inches
import  subprocess
from datetime import  datetime
import paramiko
import time
import xlsxwriter



data = [
    {"host":"10.161.4.12","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.13","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.10","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.11","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.14","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.35","uname":"root","password":"stt123"},
    {"host":"10.161.4.36","uname":"root","password":"stt123"},
    {"host":"10.161.4.37","uname":"root","password":"stt123"},
    {"host":"10.161.4.38","uname":"root","password":"stt123"},
    {"host":"10.161.4.39","uname":"root","password":"stt123"},
    {"host":"10.161.4.26","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.27","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.28","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.8.45","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.8.46","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.8.47","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.8.48","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.23","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.24","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.25","uname":"root","password":"P@ssw0rd123"},
    {"host":"10.161.4.129","uname":"root","password":"P@ssw0rd123"},
]

doc = Document()
doc.add_heading('Laporan harian status Cloudera  DC', level=2)
doc.add_heading('Monitoring harian status Cloudera DC', level=2)
doc.add_heading('Tanggal : '+ datetime.today().strftime("%d %h %y"), level=3)
doc.add_heading('Waktu : '+ datetime.today().strftime("%H:%M:%S \n"), level=3)
doc.add_table(rows=1, cols=1)
tables = doc.tables

path_photo = datetime.today().strftime("%d %h %y")+"/DC"
judul = ["1. bccdhcmdc01","2. bccdhcmdc02","3. bccdhmndc01","4. bccdhmndc02","5. bccdhwndc001","6. bccdhwndc005","7. bccdhwndc006","8. bccdhwndc007","9. bccdhwndc008","10. bccdhwndc009",
"11. kafkadc01","12. kafkadc02","13. kafkadc03","14. kafkapublicdc01","15. kafkapublicdc02","16. kafkapublicdc03","17. kafkapublicdc04","18. nifidc01","19. nifidc02","20. nifidc03","21. streamsetdc01"]
data_text = ["Health","Host Memory Usage","Host CPU Usage","Host Network Throughput","Disk Latency","df -kh"]
row = -1
cok = 0
laporan = "mkdir Laporan"
subprocess.Popen(laporan.split(), stdout=subprocess.PIPE,cwd= datetime.today().strftime("%d %h %y"))
workbook  = xlsxwriter.Workbook(datetime.today().strftime("%d %h %y")+'/Laporan/df-kh dc.xlsx')


for i in range(len(judul)):
    row += 1
    doc.tables[0].add_row()
    # ============================================
    p = tables[0].rows[row].cells[0].add_paragraph()
    r = p.add_run()
    r.add_text(judul[i])
    r.bold=True
    for y in range(len(data_text)):
        row += 1
        doc.tables[0].add_row()
        p = tables[0].rows[row].cells[0].add_paragraph()
        r = p.add_run()
        r.add_text("\t")
        r.add_text(data_text[y])
        if y != 5 :
            cok+=1
            p = tables[0].rows[row].cells[0].add_paragraph()
            r = p.add_run()
            r.add_text("\t")
            r.add_picture(path_photo+"/"+str(cok)+'.png',width=Inches(5.0), height=Inches(2.5))
        else:
            a= data[i]
            ssh = paramiko.SSHClient()
            ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            ssh.connect(a["host"], port=22, username=a["uname"], password=a["password"])


            stdin, stdout, stderr = ssh.exec_command('df -kh')
            time.sleep(1)
            b = stdout.readlines()
            xx = 0
            yy = 0
            worksheet = workbook.add_worksheet(judul[i])
            for j in b:
                xx+=1
                now = j.split()
                for k in range(6):
                    yy+=1
                    worksheet.write(int(xx), int(yy),now[k])
                    if yy == 6:
                        yy = 0
            ssh.close()
            doc.save(datetime.today().strftime("%d %h %y")+'/Laporan'+'/Laporan Harian Cloudera DC.docx')
workbook.close()